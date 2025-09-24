# word_report.py
"""
SEO audit + Word report generator.

Features:
- perform_seo_audit(url) -> returns dict with many checks
- capture_screenshot(url) -> returns PNG bytes (uses selenium if available) or raises
- create_word_report(audit_data, output_file, screenshot_bytes=None, ai_client=None)
    - output_file may be a path (str/Path) or a file-like object (BytesIO)
    - inserts screenshot if screenshot_bytes provided
    - optionally uses ai_client to add AI suggestions for Title / Meta Description
"""

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from urllib.parse import urljoin
import os
from io import BytesIO

# ---------- Audit Collector ----------

def perform_seo_audit(url: str) -> dict:
    """Fetch page and collect SEO signals. Returns dict."""
    try:
        if not url.startswith("http"):
            url = "https://" + url

        resp = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")

        # Basic meta
        title_tag = soup.find("title")
        title = title_tag.get_text(strip=True) if title_tag else ""
        title_len = len(title)

        meta_desc_tag = soup.find("meta", attrs={"name": "description"})
        meta_desc = meta_desc_tag["content"].strip() if meta_desc_tag and meta_desc_tag.get("content") else ""
        meta_desc_len = len(meta_desc)

        # H1
        h1_tags = [h.get_text(strip=True) for h in soup.find_all("h1")]

        # Images and missing alt
        images = soup.find_all("img")
        missing_alt_images = []
        for img in images:
            src = img.get("src") or img.get("data-src") or ""
            if not src:
                continue
            alt = img.get("alt") or ""
            if not alt.strip():
                abs_src = urljoin(resp.url, src)
                ext = os.path.splitext(abs_src.split("?")[0])[1].lower() or "unknown"
                missing_alt_images.append({"src": abs_src, "type": ext.lstrip(".")})

        total_images = len(images)

        # Social meta tags (og: / twitter:)
        og_tags = soup.find_all("meta", property=lambda v: v and v.startswith("og:"))
        twitter_tags = soup.find_all("meta", attrs={"name": lambda v: v and v.startswith("twitter:")})
        social_meta_present = bool(og_tags or twitter_tags)

        # Keyword frequency (simple)
        text = soup.get_text(separator=" ").lower()
        tokens = [w.strip(".,!?:;()[]\"'") for w in text.split() if len(w) > 3]
        freq = {}
        for t in tokens:
            freq[t] = freq.get(t, 0) + 1
        top_keywords = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:15]

        # Keywords usage (counts in title/meta)
        keywords_usage = []
        for kw, count in top_keywords[:10]:
            in_title = kw in title.lower()
            in_meta = kw in meta_desc.lower()
            keywords_usage.append({"keyword": kw, "count": count, "in_title": in_title, "in_meta": in_meta})

        # Google Analytics presence
        analytics_present = ("gtag(" in resp.text) or ("google-analytics.com" in resp.text) or ("analytics.js" in resp.text)

        # Favicon
        favicon_tag = soup.find("link", rel=lambda v: v and "icon" in v.lower())
        favicon_present = bool(favicon_tag and (favicon_tag.get("href") or favicon_tag.get("rel")))

        # Google Site Verification (Search Console)
        gsc_tag = soup.find("meta", attrs={"name": "google-site-verification"})
        gsc_present = bool(gsc_tag and gsc_tag.get("content"))

        # Redirects
        final_url = resp.url
        redirected = final_url.rstrip("/") != url.rstrip("/")

        # HTTPS / SSL
        https_used = final_url.startswith("https://")

        # Structured data
        structured_data = soup.find_all("script", type="application/ld+json")
        structured_present = len(structured_data) > 0

        # Custom 404 detection (simple)
        try:
            test404 = requests.get(urljoin(final_url, "/this-page-should-404-xyz"), timeout=10, headers={"User-Agent": "Mozilla/5.0"})
            custom_404 = test404.status_code == 404
        except Exception:
            custom_404 = False

        # Robots meta noindex/nofollow
        robots_meta = soup.find("meta", attrs={"name": "robots"})
        robots_content = robots_meta.get("content", "").lower() if robots_meta and robots_meta.get("content") else ""
        noindex = "noindex" in robots_content
        nofollow = "nofollow" in robots_content

        result = {
            "URL": url,
            "Final URL": final_url,
            "Redirected": redirected,
            "Title": title or "❌ Missing",
            "Title Length": title_len,
            "Meta Description": meta_desc or "❌ Missing",
            "Meta Description Length": meta_desc_len,
            "H1 Headings": h1_tags or ["❌ Missing"],
            "Missing Alt Images": missing_alt_images,
            "Total Images": total_images,
            "Social Meta Tags": "✅ Present" if social_meta_present else "❌ Missing",
            "Top Keywords": top_keywords,
            "Keywords Usage": keywords_usage,
            "Google Analytics": "✅ Found" if analytics_present else "❌ Missing",
            "Favicon": "✅ Present" if favicon_present else "❌ Missing",
            "Google Search Console": "✅ Verified" if gsc_present else "❌ Missing",
            "HTTPS": "✅ Secure" if https_used else "❌ Not Secure",
            "Structured Data": "✅ Found" if structured_present else "❌ Missing",
            "Custom 404 Page": "✅ Exists" if custom_404 else "❌ Not Found",
            "Noindex Tag": "✅ Present" if noindex else "❌ Not Present",
            "Nofollow Tag": "✅ Present" if nofollow else "❌ Not Present",
        }

        return result

    except requests.exceptions.RequestException as e:
        return {"Error": f"Could not access the website: {e}"}


# ---------- Screenshot (optional) ----------
def capture_screenshot(url: str, width: int = 1366, height: int = 900, timeout: int = 8) -> bytes:
    """
    Capture screenshot PNG bytes using selenium + webdriver-manager if available.
    Returns PNG bytes. Raises Exception if screenshot can't be taken.
    This function imports selenium dynamically to avoid top-level import errors.
    """
    # Dynamic imports to avoid hard dependency if user doesn't have selenium
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        from webdriver_manager.chrome import ChromeDriverManager
    except Exception as e:
        raise RuntimeError("Selenium or webdriver-manager not available: " + str(e))

    options = Options()
    options.add_argument("--headless=new")
    options.add_argument(f"--window-size={width},{height}")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    # turn off images for speed? we keep images to capture visual
    try:
        driver = webdriver.Chrome(ChromeDriverManager().install(), options=options)
    except TypeError:
        # webdriver-manager API difference: try Service
        from selenium.webdriver.chrome.service import Service as ChromeService
        driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)

    try:
        driver.set_page_load_timeout(timeout)
        driver.get(url)
        # try expand full height for full page screenshot
        total_height = driver.execute_script("return Math.max(document.body.scrollHeight, document.documentElement.scrollHeight);")
        driver.set_window_size(width, total_height if total_height > height else height)
        png = driver.get_screenshot_as_png()
        return png
    finally:
        try:
            driver.quit()
        except Exception:
            pass


# ---------- Word Report Generator ----------

def _ai_suggestion_text(ai_client, section_name: str, value: str) -> str:
    """
    Ask AI (Gemini) for Issue/Suggestion/Benefits. Expects ai_client to support:
      ai_client.start_chat(...).send_message(...) OR ai_client.generate(...) depending on library.
    We will attempt a couple of call patterns. If all fail, return an explanatory string.
    """
    if not ai_client:
        return ""
    prompt = f"""Provide an SEO analysis for this {section_name} value:

Value:
{value}

Format as:
Issue:
Suggestion:
Benefits for Users:
"""

    # Try common interface patterns for genai
    try:
        # genai.GenerativeModel.start_chat -> chat.send_message
        chat = ai_client.start_chat()
        resp = chat.send_message(prompt, stream=False)
        return resp.text
    except Exception:
        pass
    try:
        # genai.GenerativeModel.generate style (some libs)
        resp = ai_client.generate(prompt)
        # resp may be object; try str
        return getattr(resp, "text", str(resp))
    except Exception:
        pass
    # If all attempts fail, return nothing
    return ""


def create_word_report(audit_data: dict, output_file, screenshot_bytes: bytes = None, ai_client=None):
    """
    Create Word report and write to output_file (path or file-like).
    If screenshot_bytes provided, it will be inserted at the top.
    If ai_client provided, AI suggestions will be added for Title and Meta Description.
    """
    doc = Document()
    doc.add_heading("SEO Audit Report", level=0)
    doc.add_paragraph(f"URL: {audit_data.get('URL', 'N/A')}")
    doc.add_paragraph(f"Analyzed URL (final): {audit_data.get('Final URL', '')}")

    # Insert screenshot if available
    if screenshot_bytes:
        try:
            doc.add_heading("Website Preview", level=1)
            img_stream = BytesIO(screenshot_bytes)
            doc.add_picture(img_stream, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"[Screenshot could not be inserted: {e}]")

    # If screenshot not available, note it
    if not screenshot_bytes:
        doc.add_paragraph("Note: Screenshot was not captured (Selenium may not be available or it failed).")

    # If audit_data has Error
    if "Error" in audit_data:
        doc.add_heading("Error", level=1)
        doc.add_paragraph(audit_data["Error"])
        # Save and return
        if hasattr(output_file, "write"):
            doc.save(output_file)
        else:
            doc.save(str(output_file))
        return output_file

    # ---------- Meta Title Section (table) ----------
    doc.add_heading("Meta Title Audit", level=1)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Light Grid Accent 1"
    t.cell(0, 0).text = "Title"
    t.cell(0, 1).text = audit_data.get("Title", "")
    t.cell(1, 0).text = "Length"
    t.cell(1, 1).text = str(audit_data.get("Title Length", 0))
    t.cell(2, 0).text = "Issue"
    title_len = audit_data.get("Title Length", 0)
    if not audit_data.get("Title") or audit_data.get("Title") == "❌ Missing":
        t.cell(2, 1).text = "❌ Missing title tag."
    elif title_len < 50:
        t.cell(2, 1).text = "⚠️ Title is too short (<50 chars)."
    elif title_len > 70:
        t.cell(2, 1).text = "⚠️ Title is long (>70 chars)."
    else:
        t.cell(2, 1).text = "✅ Looks good."
    t.cell(3, 0).text = "Recommendation"
    t.cell(3, 1).text = "Keep title ~50-60 chars, include primary keyword, and make unique per page."

    # AI suggestion for Title (optional)
    if ai_client and audit_data.get("Title"):
        ai_text = _ai_suggestion_text(ai_client, "Meta Title", audit_data.get("Title", ""))
        if ai_text:
            doc.add_paragraph()
            doc.add_heading("AI - Title Analysis", level=2)
            doc.add_paragraph(ai_text)

    # ---------- Meta Description Section ----------
    doc.add_heading("Meta Description Audit", level=1)
    t = doc.add_table(rows=4, cols=2)
    t.style = "Light Grid Accent 1"
    t.cell(0, 0).text = "Meta Description"
    t.cell(0, 1).text = audit_data.get("Meta Description", "")
    t.cell(1, 0).text = "Length"
    t.cell(1, 1).text = str(audit_data.get("Meta Description Length", 0))
    t.cell(2, 0).text = "Issue"
    md_len = audit_data.get("Meta Description Length", 0)
    md_val = audit_data.get("Meta Description", "")
    if not md_val or md_val == "❌ Missing":
        t.cell(2, 1).text = "❌ Missing meta description."
    elif md_len < 120:
        t.cell(2, 1).text = "⚠️ Description is short (<120 chars)."
    elif md_len > 320:
        t.cell(2, 1).text = "⚠️ Description is long (>320 chars)."
    else:
        t.cell(2, 1).text = "✅ Looks good."
    t.cell(3, 0).text = "Recommendation"
    t.cell(3, 1).text = "Keep description 140-160 chars where possible; be compelling and include keywords."

    # AI suggestion for Meta Description (optional)
    if ai_client and md_val and md_val != "❌ Missing":
        ai_text = _ai_suggestion_text(ai_client, "Meta Description", md_val)
        if ai_text:
            doc.add_paragraph()
            doc.add_heading("AI - Meta Description Analysis", level=2)
            doc.add_paragraph(ai_text)

    # ---------- Missing Image Alt Table ----------
    doc.add_heading("Missing Image ALT Text", level=1)
    missing = audit_data.get("Missing Alt Images", [])
    if missing:
        doc.add_paragraph(f"Found {len(missing)} images missing alt text. See list below:")
        img_table = doc.add_table(rows=1, cols=3)
        img_table.style = "Medium Shading 1 Accent 2"
        hdr = img_table.rows[0].cells
        hdr[0].text = "Image URL / Path"
        hdr[1].text = "File Type"
        hdr[2].text = "Recommendation"
        for img in missing:
            row = img_table.add_row().cells
            row[0].text = img.get("src", "")
            row[1].text = (img.get("type") or "").upper()
            row[2].text = "Add descriptive alt text (describe the image, include keyword if relevant)."
    else:
        doc.add_paragraph("✅ No images missing alt text were detected.")

    # ---------- Keywords & Usage ----------
    doc.add_heading("Keywords — Top Terms & Usage", level=1)
    topk = audit_data.get("Top Keywords", [])[:15]
    if topk:
        k_table = doc.add_table(rows=1, cols=3)
        k_table.style = "Light List Accent 1"
        k_hdr = k_table.rows[0].cells
        k_hdr[0].text = "Keyword"
        k_hdr[1].text = "Frequency"
        k_hdr[2].text = "In Title / Meta"
        for kw, count in topk:
            row = k_table.add_row().cells
            usage = ""
            # find usage in keywords_usage if present
            ku = next((x for x in audit_data.get("Keywords Usage", []) if x.get("keyword") == kw), None)
            if ku:
                usage = f"{'Title' if ku.get('in_title') else ''}{' / ' if ku.get('in_title') and ku.get('in_meta') else ''}{'Meta' if ku.get('in_meta') else ''}"
            row[0].text = kw
            row[1].text = str(count)
            row[2].text = usage or "-"
    else:
        doc.add_paragraph("No significant keywords found.")

    # ---------- Advanced Checks Table ----------
    doc.add_heading("Advanced SEO Checks", level=1)
    adv_table = doc.add_table(rows=1, cols=4)
    adv_table.style = "Medium Grid 2 Accent 1"
    adv_hdr = adv_table.rows[0].cells
    adv_hdr[0].text = "Check"
    adv_hdr[1].text = "Status"
    adv_hdr[2].text = "Issue"
    adv_hdr[3].text = "Recommendation"

    def add_adv_row(name, status, issue, rec):
        r = adv_table.add_row().cells
        r[0].text = name
        r[1].text = status
        r[2].text = issue
        r[3].text = rec

    add_adv_row("Social Media Meta Tags Test", audit_data.get("Social Meta Tags", ""), 
                "Missing OpenGraph / Twitter tags affects social previews." if audit_data.get("Social Meta Tags", "").startswith("❌") else "OK",
                "Add og:title, og:description, og:image, twitter:card, etc.")
    add_adv_row("Google Analytics Test", audit_data.get("Google Analytics", ""), 
                "Analytics not found — cannot track user behaviour." if audit_data.get("Google Analytics", "").startswith("❌") else "OK",
                "Install Google Analytics (gtag.js) or GA4.")
    add_adv_row("Favicon Test", audit_data.get("Favicon", ""), 
                "No favicon present." if audit_data.get("Favicon", "").startswith("❌") else "OK",
                "Add favicon.ico or link rel icon entries.")
    add_adv_row("Google Search Console Test", audit_data.get("Google Search Console", ""), 
                "Site not verified in GSC." if audit_data.get("Google Search Console", "").startswith("❌") else "OK",
                "Verify site in Google Search Console and submit sitemap.")
    add_adv_row("URL Redirects Test", "Redirected" if audit_data.get("Redirected") else "No Redirect",
                "No canonical redirect to preferred domain." if not audit_data.get("Redirected") else "OK",
                "Ensure all variants redirect to canonical (www / non-www / http -> https).")
    add_adv_row("SSL Checker and HTTPS Test", audit_data.get("HTTPS", ""),
                "Site not fully HTTPS / SSL issues." if audit_data.get("HTTPS", "").startswith("❌") else "OK",
                "Install valid SSL certificate and force HTTPS.")
    add_adv_row("Structured Data Test", audit_data.get("Structured Data", ""),
                "No structured data found." if audit_data.get("Structured Data", "").startswith("❌") else "OK",
                "Add JSON-LD schema markup for Organization, Breadcrumb, Website, Articles etc.")
    add_adv_row("Custom 404 Error Page Test", audit_data.get("Custom 404 Page", ""),
                "No custom 404 found." if audit_data.get("Custom 404 Page", "").startswith("❌") else "OK",
                "Create a helpful branded 404 page that returns 404 status.")
    add_adv_row("Noindex Tag Test", audit_data.get("Noindex Tag", ""),
                "Noindex present — page is hidden from search engines." if audit_data.get("Noindex Tag", "").startswith("✅") else "OK",
                "Remove noindex from pages that should be indexed.")
    add_adv_row("Nofollow Tag Test", audit_data.get("Nofollow Tag", ""),
                "Nofollow present — external links will not pass link equity." if audit_data.get("Nofollow Tag", "").startswith("✅") else "OK",
                "Use nofollow sparingly only where necessary.")

    # ---------- Footer / End ----------
    doc.add_paragraph("\nEnd of report.")

    # Save to output
    if hasattr(output_file, "write"):
        # file-like object
        doc.save(output_file)
    else:
        doc.save(str(output_file))

    return output_file