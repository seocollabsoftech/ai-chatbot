# import streamlit as st
# import os
# import google.generativeai as genai
# from dotenv import load_dotenv
# import logging
# import json
# json.dump(st.session_state.messages, open("chat_history.json", "w"))

# st.markdown(
#     """
#     <style>
#         .stApp { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
#     </style>
#     """,
#     unsafe_allow_html=True
# )
# with st.chat_message("user", avatar="👤"):
#     st.markdown(prompt)
# with st.chat_message("assistant", avatar="🤖"):
#     st.markdown(ai_reply)
# theme = st.sidebar.selectbox("Theme", ["Light", "Dark"])
# if theme == "Dark":
#     st.markdown('<style> .stApp { background-color: #0e1117; color: white; } </style>', unsafe_allow_html=True)
# # Suppress ALTS warnings (from your previous fix)
# logging.getLogger('google.auth').setLevel(logging.ERROR)
# logging.getLogger('google.api_core').setLevel(logging.ERROR)

# # Load environment variables
# load_dotenv()

# # Set up Gemini client
# api_key = os.getenv("GEMINI_API_KEY")
# if not api_key:
#     st.error("Error: Add your Gemini API key to .env file!")
#     st.stop()
# genai.configure(api_key=api_key)

# # Initialize the model
# model = genai.GenerativeModel('gemini-1.5-flash')

# # Streamlit page config for beautiful layout
# st.set_page_config(
#     page_title="Gemini AI Chatbot",
#     page_icon="🤖",
#     layout="wide",  # Full-width layout
#     initial_sidebar_state="expanded"
# )

# # Custom CSS for beautiful design (add gradients, shadows, etc.)
# st.markdown("""
#     <style>
#     .stChatMessage { background-color: #f0f2f6; border-radius: 10px; padding: 10px; margin: 5px 0; }
#     .user { background-color: #007bff; color: white; }
#     .ai { background-color: #e9ecef; color: black; }
#     .main-footer { text-align: center; padding: 20px; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); color: white; }
#     </style>
# """, unsafe_allow_html=True)

# # Sidebar for settings (beautiful collapsible)
# with st.sidebar:
#     st.title("🤖 Chat Settings")
#     st.markdown("---")
#     model_choice = st.selectbox("Model", ["gemini-1.5-flash", "gemini-1.5-pro"])
#     if st.button("Clear Chat", type="secondary"):
#         st.session_state.messages = []
#     st.markdown("---")
#     st.caption("Powered by Google Gemini")

# # Initialize chat history in session state
# if "messages" not in st.session_state:
#     st.session_state.messages = []

# # Display chat history
# for message in st.session_state.messages:
#     with st.chat_message(message["role"]):
#         st.markdown(message["content"])

# # Chat input
# if prompt := st.chat_input("Type your message here..."):
#     # Add user message
#     st.session_state.messages.append({"role": "user", "content": prompt})
#     with st.chat_message("user"):
#         st.markdown(prompt)

#     # Generate AI response
#     with st.chat_message("assistant"):
#         try:
#             # Start a new chat if no history, or continue
#             if not st.session_state.messages:
#                 chat = model.start_chat()
#             else:
#                 # For simplicity, recreate chat with history (Gemini supports up to 1M tokens)
#                 chat_history = [{"role": msg["role"], "parts": [msg["content"]]} for msg in st.session_state.messages[:-1]]  # Exclude last user msg
#                 chat = model.start_chat(history=chat_history)
            
#             response = chat.send_message(prompt, stream=False)
#             ai_reply = response.text
#             st.markdown(ai_reply)
            
#             # Add AI response to history
#             st.session_state.messages.append({"role": "assistant", "content": ai_reply})
#         except Exception as e:
#             st.error(f"Oops! Error: {e}")

# # Footer for polish
# st.markdown('<div class="main-footer">© 2025 Your AI Chatbot App</div>', ugenai


import streamlit as st
import google.generativeai as genai
import os
from dotenv import load_dotenv
import logging

import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO

# --- SEO Audit Logic ---
def perform_seo_audit(url):
    """Fetches and analyzes a website's content for basic SEO elements."""
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()  # Check for bad HTTP status
        soup = BeautifulSoup(response.text, 'html.parser')

        # Collect SEO data
        title = soup.find('title')
        meta_description = soup.find('meta', attrs={'name': 'description'})
        h1_tags = [h1.get_text().strip() for h1 in soup.find_all('h1')]
        images_without_alt = [img['src'] for img in soup.find_all('img') if 'alt' not in img or not img['alt'].strip()]

        report = {
            "URL": url,
            "Title": title.string.strip() if title else "❌ Missing",
            "Meta Description": meta_description['content'].strip() if meta_description and 'content' in meta_description else "❌ Missing",
            "H1 Headings": h1_tags if h1_tags else ["❌ Missing"],
            "Images without Alt Text": images_without_alt,
            "Total Images": len(soup.find_all('img')),
        }
        return report

    except requests.exceptions.RequestException as e:
        return {"Error": f"Could not access the website: {e}"}

def create_word_report(report_data):
    """Generates a Word document from the SEO audit data."""
    doc = Document()
    doc.add_heading(f"SEO Audit Report for: {report_data.get('URL', 'N/A')}", level=1)
    
    if "Error" in report_data:
        doc.add_paragraph(f"Audit failed: {report_data['Error']}")
        return doc
    
    doc.add_paragraph("This report provides a basic analysis of the website's on-page SEO factors.")
    doc.add_heading("On-Page Analysis", level=2)
    
    # Page Title
    doc.add_heading("Title Tag", level=3)
    doc.add_paragraph(f"Title: {report_data['Title']}")
    doc.add_paragraph(f"Length: {len(report_data['Title'])} characters" if isinstance(report_data['Title'], str) else "Length: 0 characters")
    
    # Meta Description
    doc.add_heading("Meta Description", level=3)
    doc.add_paragraph(f"Description: {report_data['Meta Description']}")
    doc.add_paragraph(f"Length: {len(report_data['Meta Description'])} characters" if isinstance(report_data['Meta Description'], str) else "Length: 0 characters")

    # H1 Headings
    doc.add_heading("H1 Headings", level=3)
    for h1 in report_data['H1 Headings']:
        doc.add_paragraph(f"- {h1}", style='List Bullet')

    # Images with Missing Alt Text
    doc.add_heading("Images", level=3)
    doc.add_paragraph(f"Total Images: {report_data['Total Images']}")
    doc.add_paragraph(f"Images without Alt Text: {len(report_data['Images without Alt Text'])}")
    if report_data['Images without Alt Text']:
        doc.add_paragraph("The following images are missing alt text:")
        for img in report_data['Images without Alt Text']:
            doc.add_paragraph(f"- {img}", style='List Bullet')
    
    return doc

# --- Streamlit UI ---
st.set_page_config(
    page_title="AI Chatbot & SEO Auditor",
    page_icon="🤖",
    layout="wide"
)

# Custom CSS for a beautiful look
st.markdown("""
    <style>
    .stApp { background: linear-gradient(135deg, #1fa2ff 0%, #12d8fa 100%); }
    .stTextInput>div>div>input {
        background-color: #fff;
        border-radius: 10px;
        border: 1px solid #ccc;
        padding: 10px;
    }
    .stButton>button {
        background-color: #007bff;
        color: white;
        border-radius: 10px;
        border: none;
        padding: 10px 20px;
    }
    .stExpander {
        background-color: rgba(255, 255, 255, 0.9);
        border-radius: 10px;
    }
    </style>
""", unsafe_allow_html=True)

st.title("AI Chatbot with SEO Audit")
st.markdown("Enter a website URL below to get an on-page SEO audit report in a Word document.")

# Input for URL
url_input = st.text_input("Website URL", placeholder="e.g., https://streamlit.io")

# Button to trigger the audit
if st.button("Generate SEO Report"):
    if url_input:
        with st.spinner("Analyzing website... This may take a moment."):
            seo_data = perform_seo_audit(url_input)
        
        if "Error" in seo_data:
            st.error(seo_data["Error"])
        else:
            st.success("Audit complete! Report generated.")
            
            # Use an expander to show a summary of the findings
            with st.expander("Show Audit Summary"):
                st.write("Here is a quick overview of the findings:")
                for key, value in seo_data.items():
                    if isinstance(value, list) and value:
                        st.markdown(f"**{key}**:")
                        for item in value:
                            st.markdown(f"- {item}")
                    else:
                        st.markdown(f"**{key}**: {value}")

            # Generate and download the Word document
            word_doc = create_word_report(seo_data)
            
            # Save the document to a byte stream to make it downloadable
            doc_stream = BytesIO()
            word_doc.save(doc_stream)
            doc_stream.seek(0)
            
            # Create a download button
            st.download_button(
                label="Download SEO Report (Word)",
                data=doc_stream,
                file_name=f"seo_report_{url_input.split('//')[-1].split('/')[0]}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Please enter a URL to start the audit.")

# Suppress ALTS warnings
logging.getLogger('google.auth').setLevel(logging.ERROR)
logging.getLogger('google.api_core').setLevel(logging.ERROR)

# Load environment variables (local .env or Streamlit Cloud secrets)
load_dotenv()

# Set up Gemini client
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    st.error("Error: Add your Gemini API key to .env file or Streamlit Cloud secrets!")
    st.stop()
genai.configure(api_key=api_key)

# Initialize the model
model = genai.GenerativeModel('gemini-1.5-flash')

# Streamlit page config for beautiful layout
st.set_page_config(
    page_title="Gemini AI Chatbot",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for beautiful design
st.markdown("""
    <style>
    .stChatMessage { background-color: #f0f2f6; border-radius: 10px; padding: 10px; margin: 5px 0; }
    .user { background-color: #007bff; color: white; }
    .ai { background-color: #e9ecef; color: black; }
    .main-footer { text-align: center; padding: 20px; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); color: white; }
    </style>
""", unsafe_allow_html=True)

# Theme toggle in sidebar
with st.sidebar:
    st.title("🤖 Chat Settings")
    st.markdown("---")
    model_choice = st.selectbox("Model", ["gemini-1.5-flash", "gemini-1.5-pro"])
    theme = st.selectbox("Theme", ["Light", "Dark"])
    if theme == "Dark":
        st.markdown('<style> .stApp { background-color: #0e1117; color: white; } </style>', unsafe_allow_html=True)
    else:
        st.markdown('<style> .stApp { background: linear-gradient(135deg, #a8e063 0%, #56ab2f 100%); } </style>', unsafe_allow_html=True)
    if st.button("Clear Chat", type="secondary"):
        st.session_state.messages = []
    st.markdown("---")
    st.caption("Powered by Google Gemini")

# Initialize chat history in session state
if "messages" not in st.session_state:
    st.session_state.messages = []

# Display chat history
for message in st.session_state.messages:
    with st.chat_message(message["role"], avatar="👤" if message["role"] == "user" else "🤖"):
        st.markdown(message["content"])

# Chat input
if prompt := st.chat_input("Type your message here..."):
    # Add user message
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user", avatar="👤"):
        st.markdown(prompt)

    # Generate AI response
    with st.chat_message("assistant", avatar="🤖"):
        try:
            # Use selected model
            model = genai.GenerativeModel(model_choice)
            # Recreate chat with history
            chat_history = [{"role": msg["role"], "parts": [{"text": msg["content"]}]} for msg in st.session_state.messages[:-1]]
            chat = model.start_chat(history=chat_history)
            response = chat.send_message(prompt, stream=False)
            ai_reply = response.text
            st.markdown(ai_reply)
            
            # Add AI response to history
            st.session_state.messages.append({"role": "assistant", "content": ai_reply})
        except Exception as e:
            st.error(f"Oops! Error: {e}")

# Footer
st.markdown('<div class="main-footer">© 2025 Your AI Chatbot App</div>', unsafe_allow_html=True)